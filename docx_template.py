from io import BytesIO
from typing import Dict, Iterable, Optional, Sequence, Tuple

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from lessonplan_bot import normalize_table_rows

# Layout constants chosen to keep DOCX frame stable and close to the PDF proportions.
TABLE_LEFT_INDENT_TWIPS = 360
TABLE_WIDTH_TWIPS = 9000
DEFAULT_CELL_MARGINS_TWIPS = (96, 96, 96, 96)  # top, right, bottom, left

HEADER_INFO_COL_WIDTHS = [4500, 4500]
MATERIALS_COL_WIDTHS = [1900, 7100]
TOPIC_COL_WIDTHS = [9000]
PLAN_COL_WIDTHS = [1500, 1100, 4700, 1700]
REPORT_COL_WIDTHS = [1900, 7100]


def _safe_text(value: str, fallback: str = "") -> str:
    text = str(value or "").replace("\x00", " ").strip()
    return text or fallback


def _set_cell_text(cell, text: str, *, bold: bool = False, align_center: bool = False, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(_safe_text(text))
    run.bold = bold
    run.font.size = Pt(size)
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_table_left_indent(table, twips: int = 360) -> None:
    """Indent tables slightly so the DOCX layout better matches the PDF spacing."""
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(twips))
    tbl_ind.set(qn("w:type"), "dxa")


def _set_row_height(row, twips: int) -> None:
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    row.height = twips


def _set_table_width(table, twips: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(twips))
    tbl_w.set(qn("w:type"), "dxa")


def _set_table_layout_fixed(table) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_cell_margins(table, margins: Tuple[int, int, int, int] = DEFAULT_CELL_MARGINS_TWIPS) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)

    for tag, value in zip(("top", "right", "bottom", "left"), margins):
        mar = cell_mar.find(qn(f"w:{tag}"))
        if mar is None:
            mar = OxmlElement(f"w:{tag}")
            cell_mar.append(mar)
        mar.set(qn("w:w"), str(value))
        mar.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def _apply_col_widths_to_new_rows(table, col_widths: Sequence[int], *, start_row: int = 0) -> None:
    if not col_widths:
        return

    _set_table_layout_fixed(table)
    _set_table_width(table, sum(col_widths))
    _set_cell_margins(table)

    for row in table.rows[start_row:]:
        for idx, width in enumerate(col_widths):
            if idx < len(row.cells):
                _set_cell_width(row.cells[idx], width)


def render_week_docx(fields: Dict) -> bytes:
    document = Document()

    title = document.add_paragraph()
    title_run = title.add_run(_safe_text(fields.get("doc_title"), "주간 수업 계획서 및 보고서"))
    title_run.bold = True
    title_run.font.size = Pt(20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_table = document.add_table(rows=1, cols=2)
    info_table.style = "Table Grid"
    _set_table_left_indent(info_table)
    _apply_col_widths_to_new_rows(info_table, HEADER_INFO_COL_WIDTHS)
    left = (
        f"교사: {_safe_text(fields.get('teacher_name'), '고영찬')}\n"
        f"수업: {_safe_text(fields.get('class_name') or fields.get('subject'))}"
    )
    right = (
        f"수업날짜: {_safe_text(fields.get('lesson_datetime') or fields.get('week_label'))}\n"
        f"대상: {_safe_text(fields.get('target_group') or fields.get('class_name'))}"
    )
    _set_cell_text(info_table.rows[0].cells[0], left)
    _set_cell_text(info_table.rows[0].cells[1], right)

    materials_table = document.add_table(rows=1, cols=2)
    materials_table.style = "Table Grid"
    _set_table_left_indent(materials_table)
    _apply_col_widths_to_new_rows(materials_table, MATERIALS_COL_WIDTHS)
    _set_cell_text(materials_table.rows[0].cells[0], "수업 필요 물품 / 준비물:", bold=True)
    _set_cell_text(materials_table.rows[0].cells[1], fields.get("materials", ""))

    document.add_paragraph()

    section1 = document.add_paragraph()
    section1_run = section1.add_run("수업 주제 및 수업 목적")
    section1_run.bold = True
    section1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    topic_table = document.add_table(rows=1, cols=1)
    topic_table.style = "Table Grid"
    _set_table_left_indent(topic_table)
    _apply_col_widths_to_new_rows(topic_table, TOPIC_COL_WIDTHS)
    topic_objective = (
        f"수업 주제: {_safe_text(fields.get('lesson_topic'))}\n"
        f"수업 목적: {_safe_text(fields.get('theme_objective'))}"
    )
    _set_cell_text(topic_table.rows[0].cells[0], topic_objective)

    document.add_paragraph()

    section2 = document.add_paragraph()
    section2_run = section2.add_run("수업계획서")
    section2_run.bold = True
    section2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    plan_table = document.add_table(rows=1, cols=4)
    plan_table.style = "Table Grid"
    _set_table_left_indent(plan_table)
    _apply_col_widths_to_new_rows(plan_table, PLAN_COL_WIDTHS)
    headers = ["단계", "시간", "내용", "비고"]
    for idx, header in enumerate(headers):
        _set_cell_text(plan_table.rows[0].cells[idx], header, bold=True, align_center=True, size=11)

    rows = normalize_table_rows(fields.get("lesson_rows")) or [
        {"phase": "도입", "time": "10분", "content": "복습 및 동기 유발", "remarks": ""},
        {"phase": "전개", "time": "30분", "content": "핵심 개념 및 활동", "remarks": ""},
        {"phase": "정리", "time": "10분", "content": "형성평가 및 과제", "remarks": ""},
    ]

    for row in rows:
        cells = plan_table.add_row().cells
        _set_cell_text(cells[0], row.get("phase", ""), bold=True, align_center=True)
        _set_cell_text(cells[1], row.get("time", ""), align_center=True)
        _set_cell_text(cells[2], row.get("content", ""))
        _set_cell_text(cells[3], row.get("remarks", ""))

    _apply_col_widths_to_new_rows(plan_table, PLAN_COL_WIDTHS, start_row=1)

    for body_row in plan_table.rows[1:]:
        _set_row_height(body_row, 760)

    document.add_paragraph()

    section3 = document.add_paragraph()
    section3_run = section3.add_run("수업보고서")
    section3_run.bold = True
    section3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    report_table = document.add_table(rows=3, cols=2)
    report_table.style = "Table Grid"
    _set_table_left_indent(report_table)
    _apply_col_widths_to_new_rows(report_table, REPORT_COL_WIDTHS)

    teacher_note = _safe_text(fields.get("teacher_notes"), "특이사항 없음")
    edited_draft = _safe_text(fields.get("edited_draft"))
    if edited_draft:
        teacher_note = f"{teacher_note}\n\n[초안]\n{edited_draft}".strip()

    report_rows = [
        ("수업 평가:", _safe_text(fields.get("evaluation"), "특이사항 없음")),
        ("학생 특이 사항", _safe_text(fields.get("student_notes"), "특이사항 없음")),
        ("교사 메모", teacher_note),
    ]
    for idx, (label, value) in enumerate(report_rows):
        _set_cell_text(report_table.rows[idx].cells[0], label, bold=True)
        _set_cell_text(report_table.rows[idx].cells[1], value)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
